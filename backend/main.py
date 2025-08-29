# -*- coding: utf-8 -*-
from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Form
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import os
import io
import asyncio
import edge_tts
from PyPDF2 import PdfReader
from docx import Document
from ebooklib import epub
from bs4 import BeautifulSoup
import traceback
import uuid
import re
import unicodedata
import subprocess
from pathlib import Path
from tempfile import NamedTemporaryFile
import nest_asyncio
import aiohttp
from num2words import num2words
import chardet
import html2text
import json
from typing import Optional
import logging

# Configurar logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

nest_asyncio.apply()

app = FastAPI(title="Audiobook Generator API", version="1.0.0")

# Adicionar CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # Em produção, especifique as origens permitidas
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# === GLOBAIS E VARIÁVEIS DE CONTROLE ===

TAREFAS_JSON = "conversion_tasks.json"
GEMINI_API_KEY = None
FFMPEG_BIN = "ffmpeg"

# Dicionários em memória
cached_voices = {}
conversion_tasks = {}

# Garante que os diretórios necessários existam ao iniciar
os.makedirs("uploads", exist_ok=True)
os.makedirs("audiobooks", exist_ok=True)
os.makedirs("static", exist_ok=True)

def salvar_conversion_tasks():
    """Salva o estado atual do dicionário de tarefas em um arquivo JSON."""
    try:
        with open(TAREFAS_JSON, "w", encoding="utf-8") as f:
            json.dump(conversion_tasks, f, ensure_ascii=False, indent=2)
    except Exception as e:
        logger.error(f"⚠️ Erro ao salvar estado de tarefas: {e}")

def carregar_conversion_tasks():
    """Carrega o estado das tarefas a partir do arquivo JSON na inicialização."""
    global conversion_tasks
    if os.path.exists(TAREFAS_JSON):
        try:
            with open(TAREFAS_JSON, "r", encoding="utf-8") as f:
                conversion_tasks.update(json.load(f))
            logger.info(f"📁 Tarefas carregadas do arquivo {TAREFAS_JSON}")
        except json.JSONDecodeError:
            logger.warning(f"⚠️ Arquivo de tarefas {TAREFAS_JSON} está corrompido ou vazio. Iniciando com dicionário limpo.")
        except Exception as e:
            logger.warning(f"⚠️ Erro ao carregar arquivo de tarefas: {e}")

# Carrega tarefas salvas assim que a aplicação inicia
carregar_conversion_tasks()

# === MAPAS E PADRÕES PARA LIMPEZA DE TEXTO ===

ABREVIACOES_MAP = {
    'dr': 'Doutor', 'd': 'Dona', 'dra': 'Doutora',
    'sr': 'Senhor', 'sra': 'Senhora', 'srta': 'Senhorita',
    'prof': 'Professor', 'profa': 'Professora',
    'eng': 'Engenheiro', 'engª': 'Engenheira',
    'adm': 'Administrador', 'adv': 'Advogado',
    'exmo': 'Excelentíssimo', 'exma': 'Excelentíssima',
    'v.exa': 'Vossa Excelência', 'v.sa': 'Vossa Senhoria',
    'av': 'Avenida', 'r': 'Rua', 'km': 'Quilômetro',
    'etc': 'etcétera', 'ref': 'Referência',
    'pag': 'Página', 'pags': 'Páginas',
    'fl': 'Folha', 'fls': 'Folhas',
    'pe': 'Padre', 'dept': 'Departamento', 'depto': 'Departamento',
    'univ': 'Universidade', 'inst': 'Instituição',
    'est': 'Estado', 'tel': 'Telefone',
    'eua': 'Estados Unidos da América',
    'ed': 'Edição', 'ltda': 'Limitada'
}
ABREVIACOES_MAP_LOWER = {k.lower(): v for k, v in ABREVIACOES_MAP.items()}

CASOS_ESPECIAIS_RE = {
    r'\bV\.Exa\.(?=\s)': 'Vossa Excelência',
    r'\bV\.Sa\.(?=\s)': 'Vossa Senhoria',
    r'\bEngª\.(?=\s)': 'Engenheira'
}

CONVERSAO_CAPITULOS_EXTENSO_PARA_NUM = {
    'UM': '1', 'DOIS': '2', 'TRÊS': '3', 'QUATRO': '4', 'CINCO': '5',
    'SEIS': '6', 'SETE': '7', 'OITO': '8', 'NOVE': '9', 'DEZ': '10',
    'ONZE': '11', 'DOZE': '12', 'TREZE': '13', 'CATORZE': '14', 'QUINZE': '15',
    'DEZESSEIS': '16', 'DEZESSETE': '17', 'DEZOITO': '18', 'DEZENOVE': '19', 'VINTE': '20'
}

ABREVIACOES_QUE_NAO_TERMINAM_FRASE = set([
    'sr.', 'sra.', 'srta.', 'dr.', 'dra.', 'prof.', 'profa.', 'eng.', 'exmo.', 'exma.',
    'pe.', 'rev.', 'ilmo.', 'ilma.', 'gen.', 'cel.', 'maj.', 'cap.', 'ten.', 'sgt.',
    'cb.', 'sd.', 'me.', 'ms.', 'msc.', 'esp.', 'av.', 'r.', 'pç.', 'esq.', 'trav.',
    'jd.', 'pq.', 'rod.', 'km.', 'apt.', 'ap.', 'bl.', 'cj.', 'cs.', 'ed.', 'nº',
    'no.', 'uf.', 'cep.', 'est.', 'mun.', 'dist.', 'zon.', 'reg.', 'kg.', 'cm.',
    'mm.', 'lt.', 'ml.', 'mg.', 'seg.', 'min.', 'hr.', 'ltda.', 's.a.', 's/a',
    'cnpj.', 'cpf.', 'rg.', 'proc.', 'ref.', 'cod.', 'tel.', 'etc.', 'p.ex.', 'ex.',
    'i.e.', 'e.g.', 'vs.', 'cf.', 'op.cit.', 'loc.cit.', 'fl.', 'fls.', 'pag.',
    'p.', 'pp.', 'u.s.', 'e.u.a.', 'o.n.u.', 'i.b.m.', 'h.p.', 'obs.', 'att.',
    'resp.', 'publ.', 'ed.', 'doutora', 'senhora', 'senhor', 'doutor', 'professor',
    'professora', 'general'
])

SIGLA_COM_PONTOS_RE = re.compile(r'\b([A-Z]\.\s*)+$')

# === ROTAS ESTÁTICAS E PRINCIPAIS ===
app.mount("/static", StaticFiles(directory="static"), name="static")

@app.get("/", response_class=JSONResponse)
async def root():
    return {"message": "Audiobook Generator API", "version": "1.0.0"}

@app.get("/health", response_class=JSONResponse)
async def health_check():
    return {"status": "ok", "message": "Application is healthy."}

def _formatar_numeracao_capitulos(texto):
    def substituir_cap(match):
        tipo_cap = match.group(1).upper()
        numero_rom_arab = match.group(2)
        numero_extenso = match.group(3)
        titulo_opcional = match.group(4).strip() if match.group(4) else ""

        numero_final = ""
        if numero_rom_arab:
            numero_final = numero_rom_arab.upper()
        elif numero_extenso:
            num_ext_upper = numero_extenso.strip().upper()
            numero_final = CONVERSAO_CAPITULOS_EXTENSO_PARA_NUM.get(num_ext_upper, num_ext_upper)

        cabecalho = f"{tipo_cap} {numero_final}."
        if titulo_opcional:
            palavras_titulo = []
            for p in titulo_opcional.split():
                if p.isupper() and len(p) > 1:
                    palavras_titulo.append(p)
                else:
                    palavras_titulo.append(p.capitalize())
            titulo_formatado = " ".join(palavras_titulo)
            return f"\n\n{cabecalho}\n\n{titulo_formatado}"
        return f"\n\n{cabecalho}\n\n"

    padrao = re.compile(
        r'(?i)(cap[íi]tulo|cap\.?)\s+'
        r'(?:(\d+|[IVXLCDM]+)|([A-ZÇÉÊÓÃÕa-zçéêóãõ]+))'
        r'\s*[:\-.]?\s*'
        r'(?=\S)([^\n]*)?',
        re.IGNORECASE
    )
    texto = padrao.sub(substituir_cap, texto)

    padrao_extenso_titulo = re.compile(r'CAP[IÍ]TULO\s+([A-ZÇÉÊÓÃÕ]+)\s*[:\-]\s*(.+)', re.IGNORECASE)
    def substituir_extenso_com_titulo(match):
        num_ext = match.group(1).strip().upper()
        titulo = match.group(2).strip().title()
        numero = CONVERSAO_CAPITULOS_EXTENSO_PARA_NUM.get(num_ext, num_ext)
        return f"CAPÍTULO {numero}: {titulo}"

    texto = padrao_extenso_titulo.sub(substituir_extenso_com_titulo, texto)
    return texto

def _remover_numeros_pagina_isolados(texto):
    linhas = texto.splitlines()
    novas_linhas = []
    for linha in linhas:
        if re.match(r'^\s*\d+\s*$', linha):
            continue
        linha = re.sub(r'\s{3,}\d+\s*$', '', linha)
        novas_linhas.append(linha)
    return '\n'.join(novas_linhas)

def _normalizar_caixa_alta_linhas(texto):
    linhas = texto.splitlines()
    texto_final = []
    for linha in linhas:
        if not re.match(r'^\s*CAP[ÍI]TULO\s+[\w\d]+\.?\s*$', linha, re.IGNORECASE):
            if linha.isupper() and len(linha.strip()) > 3 and any(c.isalpha() for c in linha):
                palavras = []
                for p in linha.split():
                    if len(p) > 1 and p.isupper() and p.isalpha() and p not in ['I', 'A', 'E', 'O', 'U']:
                        if not (sum(1 for char in p if char in "AEIOU") > 0 and \
                                sum(1 for char in p if char not in "AEIOU") > 0 and len(p) <= 4):
                            palavras.append(p)
                            continue
                    palavras.append(p.capitalize())
                texto_final.append(" ".join(palavras))
            else:
                texto_final.append(linha)
        else:
            texto_final.append(linha)
    return "\n".join(texto_final)

def _corrigir_hifenizacao_quebras(texto):
    return re.sub(r'(\w+)-\s*\n\s*(\w+)', r'\1\2', texto)

def _remover_metadados_pdf(texto):
    return re.sub(r'^\s*[\w\d_-]+\.indd\s+\d+\s+\d{2}/\d{2}/\d{2,4}\s+\d{1,2}:\d{2}(:\d{2})?\s*([AP]M)?\s*$',
                  '', texto, flags=re.MULTILINE)

def _expandir_abreviacoes_numeros(texto):
    for abrev_re, expansao in CASOS_ESPECIAIS_RE.items():
        texto = re.sub(abrev_re, expansao, texto, flags=re.IGNORECASE)

    def replace_abrev_com_ponto(match):
        abrev_encontrada = match.group(1)
        expansao = ABREVIACOES_MAP_LOWER.get(abrev_encontrada.lower())
        return expansao if expansao else match.group(0)

    chaves_escapadas = [re.escape(k) for k in ABREVIACOES_MAP_LOWER if '.' not in k and 'ª' not in k]
    if chaves_escapadas:
        padrao_abrev_simples = r'\b(' + '|'.join(chaves_escapadas) + r')\.'
        texto = re.sub(padrao_abrev_simples, replace_abrev_com_ponto, texto, flags=re.IGNORECASE)

    def _converter_numero_match(match):
        num_str = match.group(0)
        try:
            if re.match(r'^\d{4}$', num_str) and 1900 <= int(num_str) <= 2100:
                return num_str
            if len(num_str) > 7:
                return num_str
            return num2words(int(num_str), lang='pt_BR')
        except Exception:
            return num_str

    texto = re.sub(r'\b\d+\b', _converter_numero_match, texto)

    def _converter_valor_monetario_match(match):
        valor_inteiro = match.group(1).replace('.', '')
        try:
            return f"{num2words(int(valor_inteiro), lang='pt_BR')} reais"
        except Exception:
            return match.group(0)

    texto = re.sub(r'R\$\s*(\d{1,3}(?:\.\d{3})*),(\d{2})', _converter_valor_monetario_match, texto)
    texto = re.sub(r'R\$\s*(\d+)(?:,00)?', lambda m: f"{num2words(int(m.group(1)), lang='pt_BR')} reais" if m.group(1) else m.group(0), texto)

    texto = re.sub(r'\b(\d+)\s*-\s*(\d+)\b',
                   lambda m: f"{num2words(int(m.group(1)), lang='pt_BR')} a {num2words(int(m.group(2)), lang='pt_BR')}",
                   texto)

    return texto

def _converter_ordinais_para_extenso(texto):
    def substituir_ordinal(match):
        numero = match.group(1)
        terminacao = match.group(2).lower()
        try:
            num_int = int(numero)
            if terminacao in ['o', 'º']:
                return num2words(num_int, lang='pt_BR', to='ordinal')
            elif terminacao in ['a', 'ª']:
                ordinal_masc = num2words(num_int, lang='pt_BR', to='ordinal')
                return ordinal_masc[:-1] + 'a' if ordinal_masc.endswith('o') else ordinal_masc
            return match.group(0)
        except ValueError:
            return match.group(0)

    padrao_ordinal = re.compile(r'\b(\d+)\s*([oaºª])(?!\w)', re.IGNORECASE)
    return padrao_ordinal.sub(substituir_ordinal, texto)

def formatar_texto_para_tts(texto_bruto: str) -> str:
    logger.info("⚙️ Aplicando formatações ao texto para TTS...")
    texto = texto_bruto

    texto = unicodedata.normalize('NFKC', texto)
    texto = texto.replace('\f', '\n\n').replace('*', '')
    for char in ['_', '#', '@']:
        texto = texto.replace(char, ' ')
    for char in ['(', ')', '\\', '[', ']']:
        texto = texto.replace(char, '')
    texto = re.sub(r'\{.*?\}', '', texto)

    texto = re.sub(r'[ \t]+', ' ', texto)
    texto = "\n".join([linha.strip() for linha in texto.splitlines() if linha.strip()])

    paragrafos_originais = texto.split('\n\n')
    paragrafos_processados = []
    for paragrafo_bruto in paragrafos_originais:
        paragrafo_bruto = paragrafo_bruto.strip()
        if not paragrafo_bruto:
            continue
        linhas = paragrafo_bruto.split('\n')
        buffer = ""
        for i, linha in enumerate(linhas):
            linha_strip = linha.strip()
            if not linha_strip:
                continue
            juntar = False
            if buffer:
                ultima_palavra = buffer.split()[-1].lower()
                termina_abrev = ultima_palavra in ABREVIACOES_QUE_NAO_TERMINAM_FRASE
                termina_sigla = re.search(r'\b[A-Z]\.$', buffer) is not None
                termina_ponto = re.search(r'[.!?…]$', buffer)
                if termina_abrev or termina_sigla:
                    juntar = True
                elif termina_ponto and linha_strip[0].isupper():
                    juntar = False
                elif not termina_ponto:
                    juntar = True
            if juntar:
                buffer += " " + linha_strip
            else:
                if buffer:
                    paragrafos_processados.append(buffer)
                buffer = linha_strip
        if buffer:
            paragrafos_processados.append(buffer)
    texto = '\n\n'.join(paragrafos_processados)

    texto = re.sub(r'[ \t]+', ' ', texto)
    texto = re.sub(r'(?<!\n)\n(?!\n)', ' ', texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)

    texto = _remover_metadados_pdf(texto)
    texto = _remover_numeros_pagina_isolados(texto)
    texto = _corrigir_hifenizacao_quebras(texto)
    texto = _formatar_numeracao_capitulos(texto)

    segmentos = re.split(r'([.!?…])\s*', texto)
    texto_reconstruido = ""
    buffer_segmento = ""
    for i in range(0, len(segmentos), 2):
        parte = segmentos[i]
        pontuacao = segmentos[i+1] if i+1 < len(segmentos) else ""
        segmento = (parte + pontuacao).strip()
        if not segmento:
            continue
        ultima = segmento.split()[-1].lower() if segmento else ""
        ultima_sem_ponto = ultima.rstrip('.!?…') if pontuacao else ultima
        abrev = ultima in ABREVIACOES_QUE_NAO_TERMINAM_FRASE or \
                ultima_sem_ponto in ABREVIACOES_QUE_NAO_TERMINAM_FRASE
        sigla = SIGLA_COM_PONTOS_RE.search(segmento) is not None
        quebra = not (pontuacao == '.' and (abrev or sigla))
        buffer_segmento += " " + segmento if buffer_segmento else segmento
        if quebra:
            texto_reconstruido += buffer_segmento + "\n\n"
            buffer_segmento = ""
    if buffer_segmento:
        texto_reconstruido += buffer_segmento
        if not re.search(r'[.!?…)]$', buffer_segmento):
            texto_reconstruido += "."
        texto_reconstruido += "\n\n"
    texto = texto_reconstruido.strip()

    texto = _normalizar_caixa_alta_linhas(texto)
    texto = _converter_ordinais_para_extenso(texto)
    texto = _expandir_abreviacoes_numeros(texto)

    for forma in ['Senhor', 'Senhora', 'Doutor', 'Doutora', 'Professor', 'Professora', 'Excelentíssimo', 'Excelentíssima']:
        texto = re.sub(r'\b' + re.escape(forma) + r'\.\s+([A-Z])', rf'{forma} \1', texto)
        texto = re.sub(r'\b' + re.escape(forma) + r'\.([A-Z])', rf'{forma} \1', texto)

    texto = re.sub(r'\b([A-Z])\.\s+([A-Z])', r'\1. \2', texto)
    texto = re.sub(r'\b([A-Z])\.\s+([A-Z][a-z])', r'\1. \2', texto)

    paragrafos = texto.split('\n\n')
    finais = []
    for p in paragrafos:
        p_strip = p.strip()
        if not p_strip:
            continue
        if not re.search(r'[.!?…)]$', p_strip) and \
           not re.match(r'^\s*CAP[ÍI]TULO\s+[\w\d]+\.?\s*$', p_strip.split('\n')[0].strip(), re.IGNORECASE):
            p_strip += '.'
        finais.append(p_strip)
    texto = '\n\n'.join(finais)
    texto = re.sub(r'[ \t]+', ' ', texto).strip()
    texto = re.sub(r'\n{2,}', '\n\n', texto)

    logger.info("✅ Formatação de texto para TTS concluída.")
    return texto.strip()

async def get_text_from_file(file_path: str, task_id: str):
    text = ""
    filename = os.path.basename(file_path)

    try:
        if filename.endswith('.pdf'):
            reader = PdfReader(file_path)
            total_pages = len(reader.pages)
            for i, page in enumerate(reader.pages):
                extracted = page.extract_text()
                if extracted:
                    text += extracted + "\n"
                progress = int(((i + 1) / total_pages) * 50)
                if (i + 1) % 10 == 0 or (i + 1) == total_pages: # Update progress less frequently
                    conversion_tasks[task_id].update({
                        "progress": progress,
                        "message": f"Extraindo texto de PDF (Página {i + 1}/{total_pages})..."
                    })

        elif filename.endswith('.txt'):
            file_content = open(file_path, 'rb').read()
            if not file_content:
                file_content = b''
            if isinstance(file_content, bytes) and len(file_content) > 0:
                detected_encoding = chardet.detect(file_content)['encoding']
                text = file_content.decode(detected_encoding or 'utf-8', errors='replace')
            else:
                text = ""
            conversion_tasks[task_id].update({
                "progress": 50,
                "message": "Texto de arquivo TXT lido."
            })

        elif filename.endswith('.docx'):
            doc = Document(file_path)
            total_paragraphs = len(doc.paragraphs)
            for i, paragraph in enumerate(doc.paragraphs):
                text += paragraph.text + "\n"
                progress = int(((i + 1) / total_paragraphs) * 50)
                if (i + 1) % 50 == 0 or (i + 1) == total_paragraphs: # Update progress less frequently
                    conversion_tasks[task_id].update({
                        "progress": progress,
                        "message": f"Extraindo texto de DOCX (Parágrafo {i + 1}/{total_paragraphs})..."
                    })

        elif filename.endswith('.epub'):
            text = _extrair_texto_de_epub_helper(file_path)
            conversion_tasks[task_id].update({
                "progress": 50,
                "message": "Texto de arquivo EPUB extraído."
            })

        conversion_tasks[task_id].update({
            "progress": 50,
            "message": "Extração de texto concluída."
        })
        logger.info(f"Extração concluída para {filename}. Total de caracteres: {len(text)}")
        return text.strip()

    except Exception as e:
        logger.error(f"Erro na extração de texto de {filename}: {e}")
        conversion_tasks[task_id].update({
            "status": "failed",
            "message": f"Erro na extração de texto: {str(e)}"
        })
        raise

def _extrair_texto_de_epub_helper(caminho_epub: str) -> str:
    texto_completo = ""
    try:
        book = epub.read_epub(caminho_epub)
        items = [item for item in book.get_items() if isinstance(item, epub.EpubHtml)] # Corrected type check for epub documents

        h = html2text.HTML2Text()
        h.ignore_links = True
        h.ignore_images = True
        h.ignore_emphasis = False
        h.body_width = 0

        for item in items:
            try:
                html_bytes = item.get_content()
                # Garantir que html_bytes seja do tipo bytes
                if not isinstance(html_bytes, (bytes, bytearray)):
                    try:
                        html_bytes = str(html_bytes).encode('utf-8', errors='replace')
                    except Exception:
                        html_bytes = b''

                if html_bytes:
                    encoding = chardet.detect(html_bytes).get('encoding') or 'utf-8'
                    try:
                        html = html_bytes.decode(encoding, errors='replace')
                    except Exception:
                        html = html_bytes.decode('utf-8', errors='replace')
                else:
                    html = ''

                soup = BeautifulSoup(html, 'html.parser')
                for tag in soup(['nav', 'header', 'footer', 'style', 'script', 'figure', 'figcaption', 'aside', 'link', 'meta']):
                    tag.decompose()

                corpo = soup.find('body') or soup
                if corpo:
                    texto_completo += h.handle(str(corpo)) + "\n\n"
            except Exception as e_file:
                logger.warning(f"⚠️ Erro no item EPUB '{item.id}': {e_file}")

        if not texto_completo.strip():
            logger.warning("⚠️ Nenhum conteúdo textual extraído do EPUB.")
            return ""
        return texto_completo
    except Exception as e:
        logger.error(f"❌ Erro geral ao processar EPUB '{caminho_epub}': {e}")
        return ""

# ================== FUNÇÃO PARA LISTAR VOZES DISPONÍVEIS ==================

cached_voices = {}

async def get_available_voices():
    """Busca vozes disponíveis, com cache para evitar requisições repetidas."""
    global cached_voices
    if cached_voices:
        return cached_voices

    logger.info("Buscando vozes Edge TTS disponíveis...")
    try:
        voices = await edge_tts.list_voices()
        pt_br_voices = {}
        for voice in voices:
            if voice["Locale"] == "pt-BR":
                name = voice["ShortName"].replace("pt-BR-", "")
                name = name.replace("Neural", " (Neural)")
                if voice["Gender"] == "Female":
                    name = f"{name} (Feminina)"
                elif voice["Gender"] == "Male":
                    name = f"{name} (Masculino)"
                pt_br_voices[voice["ShortName"]] = name.strip()

        ordered_voices = {}
        prioridade = ["pt-BR-ThalitaMultilingualNeural", "pt-BR-FranciscaNeural", "pt-BR-AntonioNeural"]
        for v in prioridade:
            if v in pt_br_voices:
                ordered_voices[v] = pt_br_voices.pop(v)
        ordered_voices.update(pt_br_voices)

        cached_voices = ordered_voices
        logger.info(f"Vozes carregadas: {len(cached_voices)} opções.")
        return cached_voices
    except Exception as e:
        logger.error(f"❌ Erro ao obter vozes Edge TTS: {e}")
        return {
            "pt-BR-ThalitaMultilingualNeural": "Thalita (Feminina, Neural) - Fallback",
            "pt-BR-FranciscaNeural": "Francisca (Feminina, Neural) - Fallback",
            "pt-BR-AntonioNeural": "Antonio (Masculino, Neural) - Fallback"
        }

# Endpoint de listagem de vozes disponíveis
@app.get("/voices", response_class=JSONResponse)
async def get_voices_endpoint():
    """Endpoint que lista as vozes disponíveis para o frontend."""
    logger.info("Recebida requisição para /voices")
    voices = await get_available_voices()
    return voices

@app.post("/process_file")
async def process_file_endpoint(
    background_tasks: BackgroundTasks,
    file: UploadFile = File(...),
    voice: str = Form(...),
    use_gemini: bool = Form(False),
    book_title: Optional[str] = Form(None)
):
    logger.info(f"Recebida requisição para /process_file com arquivo: {file.filename}, voz: {voice}, use_gemini: {use_gemini}, book_title: {book_title}")
    if not file or not file.filename:
        raise HTTPException(status_code=400, detail="Arquivo inválido ou não enviado.")

    try:
        suffix = Path(file.filename).suffix
        with NamedTemporaryFile(delete=False, dir="uploads", suffix=suffix) as temp_file:
            contents = await file.read()
            temp_file.write(contents)
            temp_file_path = temp_file.name
    except Exception as e:
        logger.error(f"❌ Erro real ao salvar arquivo temporário: {e}")
        raise HTTPException(status_code=500, detail="Erro ao salvar arquivo no servidor.")

    task_id = str(uuid.uuid4())
    conversion_tasks[task_id] = {
        "status": "in_queue",
        "message": "Tarefa recebida e na fila.",
        "progress": 0,
        "file_path": None
    }

    background_tasks.add_task(
        perform_conversion_task, temp_file_path, voice, task_id, use_gemini, book_title
    )

    salvar_conversion_tasks()
    logger.info(f"✅ Tarefa {task_id} iniciada para o arquivo {file.filename}.")

    return JSONResponse({"task_id": task_id})

@app.get("/status/{task_id}", response_class=JSONResponse)
async def get_task_status(task_id: str):
    """Endpoint para verificar o status de uma tarefa de conversão."""
    logger.info(f"🔎 Verificando status da tarefa: {task_id}")
    task = conversion_tasks.get(task_id)
    if not task:
        logger.error(f"❌ Tarefa {task_id} não encontrada no dicionário.")
        raise HTTPException(status_code=404, detail="Tarefa não encontrada ou expirada.")
    return task

@app.get("/download/{task_id}")
async def download_file(task_id: str):
    """Endpoint para baixar o audiobook finalizado."""
    task = conversion_tasks.get(task_id)
    if not task:
        raise HTTPException(status_code=404, detail="Tarefa não encontrada.")

    if task.get("status") != "completed":
        raise HTTPException(status_code=409, detail=f"A tarefa ainda não foi concluída. Status: {task.get('message')}")

    file_path = task.get("file_path")
    if not file_path or not os.path.exists(file_path):
        logger.error(f"❌ Arquivo final não encontrado para a tarefa {task_id} em {file_path}")
        raise HTTPException(status_code=404, detail="O arquivo final não foi encontrado no servidor.")
    
    filename = os.path.basename(file_path)
    return FileResponse(path=file_path, filename=filename, media_type='audio/mpeg')

@app.post("/set_gemini_api_key")
async def set_gemini_api_key_endpoint(api_key: str = Form(...)):
    """Configura a chave da API do Gemini."""
    global GEMINI_API_KEY
    if not api_key:
        raise HTTPException(status_code=400, detail="Chave API não pode ser vazia.")
    GEMINI_API_KEY = api_key
    return JSONResponse({"message": "Chave API do Gemini configurada com sucesso!"})

async def enhance_text_with_gemini(text: str) -> str:
    prompt = f"""
    Dado o texto de um livro, sua tarefa é revisá-lo e formatá-lo **exclusivamente para a narrativa principal a ser lida por um sistema de Text-to-Speech (TTS) em português do Brasil.**

    **Instruções Essenciais:**
    1.  **Prioridade Total: Focar APENAS na história/conteúdo narrativo principal.**
        * **Remova COMPLETAMENTE:** prefácios, agradecimentos, índices, bibliografias, notas de rodapé extensas, cabeçalhos e rodapés de página, números de página isolados, metadados de PDF, tabelas irrelevantes, etc.
        * **Mantenha APENAS:** O título do livro (se identificável e narrativo) e o corpo da história/conteúdo principal.
    2.  **Correção de Gramática e Ortografia:** Corrija erros de português.
    3.  **Pontuação Otimizada para Leitura:** Ajuste pontuação para leitura fluida em TTS.
    4.  **Expansão de Abreviaturas Ambíguas:** "Dr." → "Doutor", "Sra." → "Senhora", "etc." → "etcétera".
    5.  **Normalização de Números:** "1" → "um", "R$ 100,00" → "cem reais", "1º" → "primeiro".
    6.  **Remoção de Elementos Visuais:** Caracteres como `*`, `_`, `[]`, links etc. devem ser eliminados.
    7.  **Fluxo Natural:** O texto deve parecer falado, sem interrupções técnicas.
    8.  **Manter Sentido Original:** Não reescreva o conteúdo, apenas melhore a leitura em voz alta.
    9.  **Capítulos:** Certifique-se de que estão formatados como "\\n\\nCAPÍTULO X\\n\\n" com separações adequadas.
    10. **Nada de Comentários da IA:** O retorno deve conter apenas o texto narrativo revisado.

    Aqui está o texto a ser melhorado:
    ---
    {text}
    ---
    """

    chat_history = [{"role": "user", "parts": [{"text": prompt}]}]
    payload = {
        "contents": chat_history,
        "generationConfig": {
            "temperature": 0.7,
            "topP": 0.95,
            "topK": 60
        }
    }

    try:
        async with aiohttp.ClientSession() as session:
            api_url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={GEMINI_API_KEY}"
            async with session.post(api_url, headers={'Content-Type': 'application/json'}, json=payload) as response:
                response.raise_for_status()
                result = await response.json()

                if (
                    result.get("candidates") and
                    result["candidates"][0].get("content") and
                    result["candidates"][0]["content"].get("parts") and
                    result["candidates"][0]["content"]["parts"][0].get("text")
                ):
                    return result["candidates"][0]["content"]["parts"][0]["text"]
                else:
                    logger.warning("API Gemini retornou uma estrutura de resposta inesperada.")
                    logger.warning(result)
                    return text

    except aiohttp.ClientResponseError as e:
        logger.error(f"Erro HTTP da API Gemini (Status: {e.status}): {e.message}")
        if e.status == 400:
            logger.warning("Verifique se a chave API é válida ou se o texto contém conteúdo impróprio/grande demais.")
        return text
    except aiohttp.ClientError as e:
        logger.error(f"Erro de conexão ao chamar a API Gemini: {e}")
        return text
    except Exception as e:
        logger.error(f"Erro inesperado na função enhance_text_with_gemini: {e}")
        logger.error(traceback.format_exc())
        return text

def _limpar_nome_arquivo(filename: str) -> str:
    """Remove caracteres inválidos para nomes de arquivo e substitui espaços por underscore."""
    cleaned_name = re.sub(r'[<>:"/\\|?*]', '', filename)
    cleaned_name = re.sub(r'[\s-]+', '_', cleaned_name)
    cleaned_name = cleaned_name.strip('_')
    return cleaned_name[:100]

def _unificar_audios_ffmpeg(lista_arquivos_temp: list, arquivo_final: str) -> bool:
    """Une arquivos de áudio temporários em um único arquivo final usando FFmpeg."""
    if not lista_arquivos_temp:
        logger.warning("⚠️ Nenhum arquivo de áudio para unificar.")
        return False

    dir_saida = os.path.dirname(arquivo_final)
    os.makedirs(dir_saida, exist_ok=True)

    nome_lista_limpo = f"filelist_{uuid.uuid4().hex}.txt"
    lista_txt_path = os.path.join(dir_saida, nome_lista_limpo)

    try:
        with open(lista_txt_path, "w", encoding='utf-8') as f_list:
            for temp_file in lista_arquivos_temp:
                safe_path = str(Path(temp_file).resolve()).replace("'", r"\'")
                f_list.write(f"file '{safe_path}'\n")

        comando = [
            FFMPEG_BIN, '-y',
            '-f', 'concat',
            '-safe', '0',
            '-i', lista_txt_path,
            '-c', 'copy',
            arquivo_final
        ]

        process = subprocess.run(comando, stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)

        if process.returncode != 0:
            logger.error(f"❌ Erro durante a unificação (código {process.returncode}):")
            logger.error(process.stderr.decode(errors='ignore'))
            return False

        logger.info(f"✅ Unificação concluída: {os.path.basename(arquivo_final)}")
        return True

    except FileNotFoundError:
        logger.error("❌ FFmpeg não encontrado. Verifique a instalação.")
        return False
    except subprocess.CalledProcessError as e:
        logger.error(f"❌ Erro no FFmpeg: {e.stderr.decode(errors='ignore')}")
        return False
    except Exception as e:
        logger.error(f"❌ Erro inesperado: {str(e)}")
        logger.error(traceback.format_exc())
        return False
    finally:
        if os.path.exists(lista_txt_path):
            try:
                os.remove(lista_txt_path)
                logger.info(f"🧹 Lista temporária removida: {os.path.basename(lista_txt_path)}")
            except Exception as e:
                logger.warning(f"⚠️ Erro ao remover lista temporária: {e}")

async def perform_conversion_task(file_path: str, voice: str, task_id: str, use_gemini_enhancement: bool = False, book_title: Optional[str] = None): # Changed type hint to Optional[str]
    temp_chunks_dir = None

    try:
        conversion_tasks[task_id].update({
            "status": "extracting",
            "message": "Iniciando extração de texto...",
            "progress": 0
        })

        text = await get_text_from_file(file_path, task_id)

        if not text:
            conversion_tasks[task_id].update({
                "status": "failed",
                "message": "Não foi possível extrair texto do arquivo."
            })
            return

        conversion_tasks[task_id].update({
            "status": "formatting",
            "message": "Formatando texto para melhor leitura TTS (Python nativo)...",
            "progress": 55
        })
        text_formatted = formatar_texto_para_tts(text)

        if use_gemini_enhancement and GEMINI_API_KEY:
            conversion_tasks[task_id].update({
                "status": "ai_enhancing",
                "message": "Revisando e melhorando texto com IA Gemini...",
                "progress": 57
            })
            logger.info(f"Iniciando melhoria de texto com IA Gemini para tarefa {task_id}...")
            try:
                gemini_enhanced_text = await enhance_text_with_gemini(text_formatted)
                if gemini_enhanced_text.strip():
                    text_formatted = gemini_enhanced_text
                    logger.info(f"Texto melhorado com Gemini para tarefa {task_id}.")
                else:
                    logger.warning(f"Gemini retornou texto vazio para tarefa {task_id}. Usando versão Python.")
            except Exception as e_gemini:
                logger.error(f"Erro ao usar Gemini para tarefa {task_id}: {e_gemini}")
        elif use_gemini_enhancement and not GEMINI_API_KEY:
            logger.warning(f"Gemini solicitado, mas sem chave API configurada para tarefa {task_id}.")

        if not text_formatted.strip():
            conversion_tasks[task_id].update({
                "status": "failed",
                "message": "Texto vazio após formatação. Nenhuma leitura possível."
            })
            return

        if book_title and book_title.strip():
            base_filename_clean = _limpar_nome_arquivo(book_title)
            original_filename_stem = _limpar_nome_arquivo(os.path.splitext(os.path.basename(file_path))[0])
            final_audio_name_base = f"{base_filename_clean}_{original_filename_stem[:20]}" if base_filename_clean else original_filename_stem
        else:
            final_audio_name_base = _limpar_nome_arquivo(os.path.splitext(os.path.basename(file_path))[0])

        audio_filename = f"{final_audio_name_base}.mp3"
        audio_filepath = os.path.join("audiobooks", audio_filename)
        conversion_tasks[task_id]["file_path"] = audio_filepath
        conversion_tasks[task_id]["total_characters"] = len(text_formatted)

        logger.info(f"Iniciando geração de áudio com Edge TTS (Voz: {voice}) para {len(text_formatted)} caracteres formatados...")
        conversion_tasks[task_id].update({
            "status": "converting",
            "message": "Convertendo texto em áudio...",
            "progress": 60
        })

        LIMITE_CARACTERES_CHUNK_TTS = 5000
        CONCURRENCY_LIMIT = 4 # Increased concurrency limit for faster TTS generation

        text_chunks = []
        current_chunk = ""
        paragraphs = text_formatted.split('\n\n')
        for p in paragraphs:
            if not p.strip():
                continue
            if len(current_chunk) + len(p) + 2 <= LIMITE_CARACTERES_CHUNK_TTS:
                current_chunk += (("\n\n" if current_chunk else "") + p)
            else:
                if current_chunk:
                    text_chunks.append(current_chunk)
                current_chunk = p
        if current_chunk:
            text_chunks.append(current_chunk)

        if not text_chunks:
            conversion_tasks[task_id].update({
                "status": "failed",
                "message": "Nenhuma parte de texto válida para conversão após divisão."
            })
            return

        total_chunks = len(text_chunks)
        semaphore = asyncio.Semaphore(CONCURRENCY_LIMIT)
        all_tts_tasks = []

        temp_chunks_dir = os.path.join("audiobooks", f"chunks_{task_id}")
        os.makedirs(temp_chunks_dir, exist_ok=True)

        async def convert_chunk_and_save_with_retry(chunk_text, voice_id, chunk_index, max_retries=3):
            chunk_temp_file = os.path.join(temp_chunks_dir, f"chunk_{chunk_index:04d}.mp3")
            for attempt in range(max_retries):
                try:
                    async with semaphore:
                        progress_tts = int(60 + (chunk_index / total_chunks) * 35)
                        conversion_tasks[task_id].update({
                            "progress": progress_tts,
                            "message": f"Gerando áudio (Parte {chunk_index + 1}/{total_chunks}, Tentativa {attempt + 1})..."
                        })

                        communicate = edge_tts.Communicate(chunk_text, voice_id)
                        await communicate.save(chunk_temp_file)

                        if os.path.exists(chunk_temp_file) and os.path.getsize(chunk_temp_file) > 0:
                            return chunk_temp_file
                        else:
                            if os.path.exists(chunk_temp_file):
                                os.remove(chunk_temp_file)
                except Exception as e_chunk:
                    logger.warning(f"⚠️ Erro ao gerar chunk {chunk_index + 1}: {e_chunk}")
                    await asyncio.sleep(2 ** attempt)
            return None

        for i, chunk_text in enumerate(text_chunks):
            task = asyncio.create_task(convert_chunk_and_save_with_retry(chunk_text, voice, i))
            all_tts_tasks.append(task)

        results = await asyncio.gather(*all_tts_tasks)
        successful_chunk_files = [res for res in results if res and os.path.exists(res)]

        if not successful_chunk_files:
            conversion_tasks[task_id].update({
                "status": "failed",
                "message": "Nenhum áudio válido foi gerado para o audiobook."
            })
            return

        conversion_tasks[task_id].update({
            "status": "merging_audio",
            "message": "Unificando partes do áudio...",
            "progress": 98
        })

        if _unificar_audios_ffmpeg(successful_chunk_files, audio_filepath):
            conversion_tasks[task_id].update({
                "status": "completed",
                "message": "Audiobook pronto para download!",
                "progress": 100
            })
        else:
            conversion_tasks[task_id].update({
                "status": "failed",
                "message": "Falha ao unificar partes do áudio. O audiobook pode estar incompleto."
            })
            return

    except Exception as e:
        logger.error(f"Erro na conversão da tarefa {task_id}: {e}")
        logger.error(traceback.format_exc())
        conversion_tasks[task_id].update({
            "status": "failed",
            "message": f"Erro na conversão: {str(e)}"
        })

    finally:
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
                logger.info(f"Arquivo de texto original temporário {os.path.basename(file_path)} removido.")
            except Exception as e:
                logger.warning(f"⚠️ Erro ao remover arquivo original temporário: {e}")

        if temp_chunks_dir and os.path.exists(temp_chunks_dir):
            try:
                for temp_chunk_file in os.listdir(temp_chunks_dir):
                    try:
                        os.remove(os.path.join(temp_chunks_dir, temp_chunk_file))
                    except Exception as e_clean:
                        logger.warning(f"⚠️ Erro ao remover chunk '{temp_chunk_file}': {e_clean}")
                os.rmdir(temp_chunks_dir)
                logger.info(f"🧹 Diretório de chunks temporários removido: {os.path.basename(temp_chunks_dir)}")
            except Exception as e_rmdir:
                logger.warning(f"⚠️ Erro ao remover diretório de chunks temporários: {e_rmdir}")


@app.post("/shutdown")
async def shutdown_application():
    """Endpoint para sinalizar shutdown criando arquivo shutdown.flag."""
    try:
        with open("../shutdown.flag", "w") as f:
            f.write("shutdown")
        logger.info("🛑 Arquivo shutdown.flag criado.")

        # Tentar executar stop-local.bat automaticamente (caso start-local.bat nao esteja monitorando)
        try:
            project_root = Path(__file__).resolve().parent.parent
            stop_script = project_root / "stop-local.bat"
            if stop_script.exists():
                logger.info(f"🛑 Tentando executar {stop_script} para encerrar serviços...")
                # Executa o script em background via cmd
                subprocess.Popen(["cmd", "/c", str(stop_script)], cwd=str(project_root), stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            else:
                logger.warning(f"Arquivo de parada nao encontrado: {stop_script}")
        except Exception as e_run:
            logger.warning(f"Erro ao tentar executar stop-local.bat automaticamente: {e_run}")

        # Tentativa adicional: encerrar processos que estejam escutando nas portas 8000 e 3000 (Windows)
        try:
            import platform
            if platform.system().lower().startswith("win"):
                logger.info("🛑 Tentando identificar processos nas portas 8000/3000 e finalizá-los (Windows)...")
                try:
                    netstat_out = subprocess.check_output("netstat -ano", shell=True, stderr=subprocess.DEVNULL)
                    netstat_text = netstat_out.decode(errors='ignore')
                    for port in (8000, 3000):
                        for line in netstat_text.splitlines():
                            if f":{port} " in line or f":{port}\t" in line or f":{port}" in line:
                                parts = line.split()
                                if parts:
                                    pid = parts[-1]
                                    if pid.isdigit():
                                        try:
                                            subprocess.run(["taskkill", "/PID", pid, "/F"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                                            logger.info(f"🛑 Finalizado PID {pid} que atendia porta {port}")
                                        except Exception as e_kill:
                                            logger.warning(f"Erro ao finalizar PID {pid}: {e_kill}")
                except Exception as e_net:
                    logger.warning(f"Erro ao executar netstat/taskkill: {e_net}")
            else:
                logger.info("🛑 Sistema nao-Windows detectado; pulando tentativa automatica de taskkill.")
        except Exception as e_extra:
            logger.warning(f"Erro na tentativa adicional de finalizacao de portas: {e_extra}")

        return JSONResponse({"message": "Sinal de desligamento enviado. Encerramento em andamento."})
    except Exception as e:
        logger.error(f"Erro ao criar arquivo shutdown.flag: {e}")
        raise HTTPException(status_code=500, detail="Erro ao sinalizar desligamento.")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
